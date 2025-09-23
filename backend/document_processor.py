import os
import openai
from typing import Union
import PyPDF2
from docx import Document
import io

class DocumentProcessor:
    def __init__(self):
        """Initialize document processor with OpenAI configuration"""
        self.openai_client = openai.OpenAI(
            api_key=os.getenv('OPENAI_API_KEY', 'your-openai-api-key-here')
        )
        self.model = os.getenv('OPENAI_MODEL', 'gpt-4')
        
        # Supported file types
        self.supported_types = {
            'application/pdf': self._extract_pdf_text,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._extract_docx_text,
            'application/msword': self._extract_doc_text,
            'text/plain': self._extract_text_text,
        }
    
    def extract_text(self, file) -> str:
        """Extract text from uploaded file based on its type"""
        try:
            content_type = file.content_type
            filename = file.filename.lower()
            
            # Determine extraction method
            if content_type in self.supported_types:
                extractor = self.supported_types[content_type]
            elif filename.endswith('.pdf'):
                extractor = self._extract_pdf_text
            elif filename.endswith('.docx'):
                extractor = self._extract_docx_text
            elif filename.endswith('.doc'):
                extractor = self._extract_doc_text
            elif filename.endswith('.txt'):
                extractor = self._extract_text_text
            else:
                raise ValueError(f"Unsupported file type: {content_type}")
            
            # Extract text
            text = extractor(file)
            
            # Clean and validate text
            cleaned_text = self._clean_text(text)
            
            if not cleaned_text.strip():
                raise ValueError("No text content could be extracted from the file")
            
            return cleaned_text
            
        except Exception as e:
            raise Exception(f"Failed to extract text from {file.filename}: {str(e)}")
    
    def _extract_pdf_text(self, file) -> str:
        """Extract text from PDF file"""
        try:
            # Read PDF content
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            
            # Extract text from each page
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            
            return text
            
        except Exception as e:
            raise Exception(f"PDF extraction failed: {str(e)}")
    
    def _extract_docx_text(self, file) -> str:
        """Extract text from DOCX file"""
        try:
            # Read DOCX content
            doc = Document(file)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + "\n"
            
            return text
            
        except Exception as e:
            raise Exception(f"DOCX extraction failed: {str(e)}")
    
    def _extract_doc_text(self, file) -> str:
        """Extract text from DOC file (basic implementation)"""
        try:
            # For .doc files, we'll use a basic approach
            # In production, you might want to use python-docx2txt or similar
            file.seek(0)
            content = file.read()
            
            # Try to extract readable text (basic approach)
            text = ""
            try:
                # Convert bytes to string and extract readable characters
                if isinstance(content, bytes):
                    content = content.decode('utf-8', errors='ignore')
                
                # Extract alphanumeric and common punctuation
                import re
                text = re.sub(r'[^\w\s\.\,\!\?\;\:\-\(\)\[\]\{\}]', ' ', content)
                text = re.sub(r'\s+', ' ', text).strip()
                
            except Exception:
                # Fallback: return raw content as string
                text = str(content)
            
            return text
            
        except Exception as e:
            raise Exception(f"DOC extraction failed: {str(e)}")
    
    def _extract_text_text(self, file) -> str:
        """Extract text from plain text file"""
        try:
            file.seek(0)
            content = file.read()
            
            if isinstance(content, bytes):
                content = content.decode('utf-8', errors='ignore')
            
            return content
            
        except Exception as e:
            raise Exception(f"Text extraction failed: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        import re
        text = re.sub(r'\s+', ' ', text)
        
        # Remove non-printable characters
        text = ''.join(char for char in text if char.isprintable() or char in '\n\t')
        
        # Normalize line endings
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # Remove excessive newlines
        text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
        
        return text.strip()
    
    def generate_summary(self, text: str, max_length: int = 500) -> str:
        """Generate AI-powered summary of the document text"""
        try:
            # Truncate text if it's too long for the API
            max_tokens = 8000  # Conservative limit for GPT-4
            if len(text) > max_tokens * 4:  # Rough estimate: 1 token ≈ 4 characters
                text = text[:max_tokens * 4]
            
            # Create prompt for summary generation
            prompt = f"""Please provide a concise summary of the following student club information. Focus on:
1. The club's mission and purpose
2. Key activities and events
3. Notable achievements or impact
4. Target audience or membership

Keep the summary to approximately {max_length} characters and write it in a professional tone suitable for business outreach.

Club Information:
{text}

Summary:"""

            # Generate summary using OpenAI
            response = self.openai_client.chat.completions.create(
                model=self.model,
                messages=[
                    {
                        "role": "system",
                        "content": "You are an AI assistant that creates concise, professional summaries of student club information for business outreach purposes."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                max_tokens=300,
                temperature=0.3,
                top_p=0.9
            )
            
            summary = response.choices[0].message.content.strip()
            
            # Ensure summary is within length limit
            if len(summary) > max_length:
                summary = summary[:max_length-3] + "..."
            
            return summary
            
        except Exception as e:
            # Fallback to basic summary if AI fails
            fallback_summary = self._generate_fallback_summary(text, max_length)
            print(f"AI summary generation failed, using fallback: {str(e)}")
            return fallback_summary
    
    def _generate_fallback_summary(self, text: str, max_length: int) -> str:
        """Generate a basic summary without AI as fallback"""
        try:
            # Simple text processing for fallback
            sentences = text.split('.')
            summary_sentences = []
            current_length = 0
            
            for sentence in sentences:
                sentence = sentence.strip()
                if not sentence:
                    continue
                
                # Add sentence if it fits
                if current_length + len(sentence) + 1 <= max_length:
                    summary_sentences.append(sentence)
                    current_length += len(sentence) + 1
                else:
                    break
            
            summary = '. '.join(summary_sentences)
            if summary and not summary.endswith('.'):
                summary += '.'
            
            return summary
            
        except Exception:
            # Last resort: return first portion of text
            return text[:max_length-3] + "..." if len(text) > max_length else text
    
    def chunk_text(self, text: str, chunk_size: int = 4000) -> list[str]:
        """Split long text into chunks for processing"""
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        current_chunk = ""
        
        # Split by sentences to maintain readability
        sentences = text.split('.')
        
        for sentence in sentences:
            sentence = sentence.strip() + '.'
            
            if len(current_chunk) + len(sentence) <= chunk_size:
                current_chunk += sentence + " "
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = sentence + " "
        
        # Add the last chunk
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks
    
    def get_file_info(self, file) -> dict:
        """Get information about the uploaded file"""
        try:
            file.seek(0, 2)  # Seek to end
            file_size = file.tell()
            file.seek(0)  # Reset to beginning
            
            return {
                'filename': file.filename,
                'content_type': file.content_type,
                'size_bytes': file_size,
                'size_mb': round(file_size / (1024 * 1024), 2),
                'supported': file.content_type in self.supported_types
            }
        except Exception as e:
            return {
                'filename': file.filename,
                'error': str(e)
            }






